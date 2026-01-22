import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import tempfile
from streamlit_mermaid import st_mermaid

# --- 1. CONFIGURATION & SETUP ---
st.set_page_config(page_title="VerbaFlow Team", page_icon="üéôÔ∏è", layout="wide")

# Custom CSS for that "Standout" UI
st.markdown("""
<style>
    .main { background-color: #0E1117; color: #FAFAFA; }
    h1, h2, h3 { color: #FF4B4B; font-family: 'Helvetica Neue', sans-serif; }
    .stButton>button {
        width: 100%; border-radius: 8px; border: none;
        background: linear-gradient(90deg, #FF4B4B 0%, #FF914D 100%);
        color: white; font-weight: bold; padding: 10px; transition: 0.3s;
    }
    .stButton>button:hover { transform: scale(1.02); box-shadow: 0 4px 15px rgba(255, 75, 75, 0.4); }
    .agent-box {
        border-left: 5px solid #FF4B4B; background-color: #1E1E1E;
        padding: 20px; border-radius: 5px; margin-bottom: 20px;
    }
    .chat-bubble {
        background-color: #262730; padding: 15px; border-radius: 15px;
        margin-bottom: 10px; border: 1px solid #444;
    }
    .team-card {
        background-color: #1E1E1E; padding: 10px; border-radius: 5px;
        margin-bottom: 10px; border: 1px solid #333; font-size: 0.9em;
    }
</style>
""", unsafe_allow_html=True)

# --- 2. THE INTELLIGENT AGENT ---
class TranscriptionAgent:
    def __init__(self, api_key):
        genai.configure(api_key=api_key)
        # Using the Free Tier capable model
        self.model = genai.GenerativeModel('gemini-flash-latest')

    def listen_and_transcribe(self, audio_file_path):
        """Native hearing with Speaker Identification."""
        try:
            audio_file = genai.upload_file(path=audio_file_path)
            prompt = """
            Listen to this audio carefully. It may contain English, Hindi, or Hinglish.
            1. Transcribe it exactly as spoken.
            2. Distinguish between speakers (e.g., 'Speaker 1:', 'Speaker 2:').
            3. If it's a lecture, label the main speaker as 'Lecturer'.
            """
            response = self.model.generate_content([prompt, audio_file])
            return response.text
        except Exception as e:
            return f"Agent Error: {str(e)}"

    def think_and_process(self, text, task):
        prompts = {
            "Summarize": f"Create a comprehensive bullet-point summary:\n\n{text}",
            "Elaborate": f"Explain the concepts simply for a beginner:\n\n{text}",
            "Action Items": f"Identify tasks and assignments as a checklist:\n\n{text}",
            "Quiz": f"Generate 3 quiz questions with answers at the bottom:\n\n{text}",
            "Mind Map": f"""
            Create a Mermaid.js flowchart code to visualize connections.
            - Start with 'graph TD'.
            - Use short node labels.
            - Do NOT use markdown code blocks. Return raw code only.
            Text: {text}
            """
        }
        prompt = prompts.get(task, "Analyze this text.")
        response = self.model.generate_content(prompt)
        return response.text.replace("```mermaid", "").replace("```", "").strip()

    def ask_question(self, transcript, question):
        prompt = f"Context: {transcript}\n\nQuestion: {question}\n\nAnswer concisely based on context:"
        response = self.model.generate_content(prompt)
        return response.text

# --- 3. FILE GENERATORS ---
def create_pdf(title, transcript, notes):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    def clean(t): return t.encode('latin-1', 'replace').decode('latin-1')
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, clean(title), ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "AI Notes:", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 10, clean(notes))
    pdf.ln(10)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Transcript:", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, clean(transcript))
    return pdf.output(dest='S').encode('latin-1')

def create_docx(transcript, notes):
    doc = Document()
    doc.add_heading('VerbaFlow Report', 0)
    doc.add_heading('AI Notes', level=1)
    doc.add_paragraph(notes)
    doc.add_heading('Transcript', level=1)
    doc.add_paragraph(transcript)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as f: return f.read()

# --- 4. THE UI ---
def main():
    # --- SIDEBAR: TEAM CREDITS ---
    with st.sidebar:
        st.title("‚öôÔ∏è Settings")
        api_key = st.text_input("Gemini API Key", type="password")
        
        st.markdown("---")
        st.subheader("CREATOR")
        
        st.markdown("""
       
        <div class="team-card">
            <b>Arvindsundaram S</b><br>
            <i>AI Logic & Backend</i>
        </div>
        
        """, unsafe_allow_html=True)

    # --- MAIN CONTENT ---
    st.markdown('<h1 style="text-align: center;">üéôÔ∏è VerbaFlow Pro</h1>', unsafe_allow_html=True)
    
    if not api_key:
        st.warning("üëà Please enter the API Key in the sidebar.")
        return

    agent = TranscriptionAgent(api_key)

    # 1. RECORDING
    st.markdown("### 1. Audio Source")
    audio_value = st.audio_input("Record Lecture")

    if audio_value:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(audio_value.read())
            tmp_path = tmp.name

        if st.button("üìù Transcribe (Identify Speakers)", use_container_width=True):
            with st.spinner("Agent is listening..."):
                transcript = agent.listen_and_transcribe(tmp_path)
                st.session_state['transcript'] = transcript
                st.rerun()

    # 2. TRANSCRIPT & ACTIONS
    if 'transcript' in st.session_state:
        st.markdown("### 2. Transcript")
        st.markdown(f"<div class='agent-box'>{st.session_state['transcript']}</div>", unsafe_allow_html=True)

        st.markdown("### 3. Knowledge Engine")
        c1, c2, c3, c4 = st.columns(4)
        task = None
        if c1.button("üìù Summarize"): task = "Summarize"
        if c2.button("üß† Explain"): task = "Elaborate"
        if c3.button("‚úÖ Action Items"): task = "Action Items"
        if c4.button("üó∫Ô∏è Mind Map"): task = "Mind Map"

        if task:
            with st.spinner(f"Agent is generating {task}..."):
                result = agent.think_and_process(st.session_state['transcript'], task)
                st.session_state['result'] = result
                st.session_state['task'] = task

        if 'result' in st.session_state:
            st.markdown(f"### Result: {st.session_state['task']}")
            if st.session_state['task'] == "Mind Map":
                try:
                    st_mermaid(st.session_state['result'], height="400px")
                except:
                    st.code(st.session_state['result'])
            else:
                st.info(st.session_state['result'])

            # Export
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.download_button("Download DOCX", create_docx(st.session_state['transcript'], st.session_state['result']), "notes.docx")
            with col_d2:
                st.download_button("Download PDF", create_pdf("VerbaFlow", st.session_state['transcript'], st.session_state['result']), "notes.pdf")

        # Chat
        st.markdown("---")
        st.markdown("### üí¨ Chat with Audio")
        user_q = st.text_input("Ask a question about the recording...")
        if user_q:
            with st.spinner("Thinking..."):
                answer = agent.ask_question(st.session_state['transcript'], user_q)
                st.markdown(f"<div class='chat-bubble'><b>Agent:</b> {answer}</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()